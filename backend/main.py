from fastapi import FastAPI, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from datetime import datetime

import models
from database import engine, SessionLocal
from ai_pipeline import run_ai_pipeline
from pydantic import BaseModel
import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from dotenv import load_dotenv
load_dotenv()

models.Base.metadata.create_all(bind=engine)

# Safely run migrations for newly added columns
try:
    from sqlalchemy import text
    with engine.connect() as conn:
        conn.execute(text("ALTER TABLE devices ADD COLUMN remainingUsefulLife REAL DEFAULT 36.0"))
        conn.commit()
except Exception:
    pass
app = FastAPI(title="DeviceGuardian AI Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

class SearchQuery(BaseModel):
    query: str

@app.post("/api/search")
def ai_search(data: SearchQuery, db: Session = Depends(get_db)):
    try:
        import google.generativeai as genai
        api_key = os.getenv("GEMINI_API_KEY")
        if api_key:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(
                f"You are DeviceGuardian AI, a device health assistant. Answer this query concisely: {data.query}"
            )
            return {"summary": response.text}
    except Exception as e:
        print(f"Gemini API search failed: {e}")
        
    return {
        "summary": f"I couldn't reach the live AI. For query '{data.query}': Keep battery below 45C and SSD usage below 85% for optimal device health."
    }

@app.get("/")
def read_root():
    return {"message": "DeviceGuardian AI Backend is running"}

@app.get("/api/accuracy")
def get_accuracy():
    try:
        import json
        with open("accuracy_proof.json", "r") as f:
            return json.load(f)
    except Exception as e:
        return {"error": str(e)}

@app.post("/devices")
def register_device(telemetry: models.TelemetryCreate, db: Session = Depends(get_db)):
    # 1. Update or create the Device in the database
    device = db.query(models.DeviceRecord).filter(models.DeviceRecord.id == telemetry.deviceId).first()
    if not device:
        device = models.DeviceRecord(
            id=telemetry.deviceId,
            name=telemetry.name,
            deviceType=telemetry.deviceType,
            status=telemetry.status
        )
        db.add(device)
    else:
        device.status = telemetry.status
        device.lastUpdated = datetime.utcnow()

    # 2. Add telemetry data point
    new_telemetry = models.TelemetryRecord(
        deviceId=telemetry.deviceId,
        cpu=telemetry.cpu,
        ram=telemetry.ram,
        battery=telemetry.battery,
        temperature=telemetry.temperature,
        ssd=telemetry.ssd
    )
    db.add(new_telemetry)
    db.commit()
    
    return {"message": "Telemetry received successfully"}

@app.delete("/devices/{device_id}")
def delete_device(device_id: str, db: Session = Depends(get_db)):
    device = db.query(models.DeviceRecord).filter(models.DeviceRecord.id == device_id).first()
    if device:
        # Delete related telemetry and predictions
        db.query(models.TelemetryRecord).filter(models.TelemetryRecord.deviceId == device_id).delete()
        db.query(models.AIPrediction).filter(models.AIPrediction.deviceId == device_id).delete()
        db.delete(device)
        db.commit()
        
    # Also delete mapping from PostgreSQL device_mappings table directly
    try:
        from sqlalchemy import text
        db.execute(text("DELETE FROM device_mappings WHERE device_uuid = :uuid"), {"uuid": device_id})
        db.commit()
    except Exception as e:
        print(f"Non-fatal: SQLite fallback or Postgres mapping delete failed: {e}")
        
    return {"message": "Device and mapping deleted successfully"}

@app.get("/devices")
def get_devices(db: Session = Depends(get_db)):
    devices = db.query(models.DeviceRecord).all()
    result = []
    for d in devices:
        # Get all telemetry history (oldest first for cycle estimation)
        all_telemetry = db.query(models.TelemetryRecord)\
            .filter(models.TelemetryRecord.deviceId == d.id)\
            .order_by(models.TelemetryRecord.id.asc())\
            .all()
        
        # Latest reading is the last element
        latest_telemetry = all_telemetry[-1] if all_telemetry else None
            
        battery = latest_telemetry.battery if latest_telemetry else 100
        temp = latest_telemetry.temperature if latest_telemetry else 35
        
        # Fetch the latest prediction from DB (calculated by predictions API)
        latest_pred = db.query(models.AIPrediction)\
            .filter(models.AIPrediction.deviceId == d.id)\
            .order_by(models.AIPrediction.id.desc())\
            .first()
            
        health = 100
        if latest_pred:
            health = latest_pred.healthScore
        elif latest_telemetry:
            telemetry_dict = {
                "deviceId": latest_telemetry.deviceId,
                "cpu": latest_telemetry.cpu,
                "ram": latest_telemetry.ram,
                "battery": latest_telemetry.battery,
                "temperature": latest_telemetry.temperature,
                "ssd": latest_telemetry.ssd
            }
            try:
                prediction = run_ai_pipeline(telemetry_dict, all_telemetry)
                health = prediction["healthScore"]
                if d.deviceType == "phone":
                    rul = d.remainingUsefulLife if d.remainingUsefulLife is not None else 36.0
                    health = int(round(80 + 20 * (rul / 36.0)))
            except Exception as e:
                print(f"Error running pipeline in list API: {e}")
            
        result.append({
            "id": d.id,
            "name": d.name,
            "type": d.deviceType,
            "status": d.status,
            "lastUpdated": d.lastUpdated.isoformat(),
            "battery": battery,
            "temperature": temp,
            "healthScore": health
        })
    return result

def sync_device_from_supabase(device_id: str, db: Session):
    try:
        import requests
        url = f"https://lonsqhuudhiffjitmcbh.supabase.co/rest/v1/telemetry?device_uuid=eq.{device_id}"
        headers = {
            "apikey": "sb_publishable_huLEhuc-J4bal6hQRkPf5w_O16MKv6V",
            "Authorization": "Bearer sb_publishable_huLEhuc-J4bal6hQRkPf5w_O16MKv6V"
        }
        r = requests.get(url, headers=headers, timeout=5)
        if r.status_code == 200 and r.json():
            data = r.json()[0]
            payload = data.get("payload", {})
            name = data.get("device_name", "Unknown Device")
            
            def safe_float(val, default=0.0):
                if val is None:
                    return default
                if isinstance(val, str):
                    val = val.replace("%", "").strip()
                try:
                    return float(val)
                except (ValueError, TypeError):
                    return default

            battery_map = payload.get("battery") if isinstance(payload.get("battery"), dict) else {}
            cpu_map = payload.get("cpu") if isinstance(payload.get("cpu"), dict) else {}
            storage_map = payload.get("storage") if isinstance(payload.get("storage"), dict) else {}
            memory_map = payload.get("memory") if isinstance(payload.get("memory"), dict) else {}

            battery_val = battery_map.get("level") if battery_map.get("level") is not None else battery_map.get("percentage")
            battery = safe_float(battery_val, 100.0)

            temp = safe_float(cpu_map.get("temperature_c"), 35.0)
            if temp == 35.0 and cpu_map.get("gpu_temperature_c") is not None:
                temp = safe_float(cpu_map.get("gpu_temperature_c"), 35.0)

            ssd_val = storage_map.get("usage_percent") if storage_map.get("usage_percent") is not None else storage_map.get("disk_usage_percent")
            ssd = safe_float(ssd_val, 50.0)

            cpu = safe_float(cpu_map.get("usage_percent"), 20.0)

            ram_val = memory_map.get("usage_percent") if memory_map.get("usage_percent") is not None else memory_map.get("ram_usage_percent")
            ram = safe_float(ram_val, 50.0)

            is_charging = battery_map.get("is_charging") if battery_map.get("is_charging") is not None else battery_map.get("charging")
            is_charging = True if is_charging in [True, "true", "True"] else False
            status = "Charging" if is_charging else "Healthy"
            
            device_type = "phone"
            system_map = payload.get("system") if isinstance(payload.get("system"), dict) else {}
            win_version = str(system_map.get("windows_version", "")).lower()
            if (
                "windows" in name.lower() or 
                "laptop" in name.lower() or 
                "pc" in name.lower() or 
                "ashwin" in name.lower() or 
                "amudieshwar" in name.lower() or
                "devesh" in name.lower() or
                "windows" in win_version or
                "gpus" in payload or
                "disk_health" in payload
            ):
                device_type = "laptop"

            # Parse actual timestamp from Supabase if available
            try:
                updated_at_str = data.get("updated_at", "")
                actual_timestamp = datetime.fromisoformat(updated_at_str.replace("Z", "+00:00")).replace(tzinfo=None)
            except Exception:
                actual_timestamp = datetime.utcnow()

            # Parse remaining useful life from health_prediction payload
            if device_type == "phone":
                rul_val = payload.get("native_remaining_useful_life_months") or payload.get("health_prediction", {}).get("remaining_useful_life_months", 36.0)
            else:
                rul_val = payload.get("health_prediction", {}).get("remaining_useful_life_months", 36.0)
            rul = safe_float(rul_val, 36.0)

            device = db.query(models.DeviceRecord).filter(models.DeviceRecord.id == device_id).first()
            if not device:
                device = models.DeviceRecord(
                    id=device_id,
                    name=name,
                    deviceType=device_type,
                    status=status,
                    lastUpdated=actual_timestamp,
                    remainingUsefulLife=rul
                )
                db.add(device)
            else:
                device.name = name
                device.deviceType = device_type
                device.status = status
                device.lastUpdated = actual_timestamp
                device.remainingUsefulLife = rul
            
            new_telemetry = models.TelemetryRecord(
                deviceId=device_id,
                cpu=cpu,
                ram=ram,
                battery=battery,
                temperature=temp,
                ssd=ssd,
                timestamp=actual_timestamp
            )
            db.add(new_telemetry)
            db.commit()
            return True
    except Exception as e:
        print(f"Supabase sync failed: {e}")
    return False

@app.get("/devices/{device_id}")
def get_device(device_id: str, db: Session = Depends(get_db)):
    # Always attempt to sync latest telemetry from Supabase
    sync_device_from_supabase(device_id, db)
    
    device = db.query(models.DeviceRecord).filter(models.DeviceRecord.id == device_id).first()
    if not device:
        raise HTTPException(status_code=404, detail="Device not found")
        
    latest_telemetry = db.query(models.TelemetryRecord)\
        .filter(models.TelemetryRecord.deviceId == device_id)\
        .order_by(models.TelemetryRecord.id.desc())\
        .first()
        
    battery = latest_telemetry.battery if latest_telemetry else 100
    temp = latest_telemetry.temperature if latest_telemetry else 35
    ssd = latest_telemetry.ssd if latest_telemetry else 50
    cpu = latest_telemetry.cpu if latest_telemetry else 0
    ram = latest_telemetry.ram if latest_telemetry else 0
    
    # Fetch the latest prediction from DB (calculated by predictions API)
    latest_pred = db.query(models.AIPrediction)\
        .filter(models.AIPrediction.deviceId == device_id)\
        .order_by(models.AIPrediction.id.desc())\
        .first()
        
    health = 100
    if latest_pred:
        health = latest_pred.healthScore
    elif latest_telemetry:
        all_telemetry = db.query(models.TelemetryRecord)\
            .filter(models.TelemetryRecord.deviceId == device_id)\
            .order_by(models.TelemetryRecord.id.asc())\
            .all()
            
        telemetry_dict = {
            "deviceId": device_id,
            "cpu": cpu,
            "ram": ram,
            "battery": battery,
            "temperature": temp,
            "ssd": ssd
        }
        try:
            prediction = run_ai_pipeline(telemetry_dict, all_telemetry)
            health = prediction["healthScore"]
            if device.deviceType == "phone":
                rul = device.remainingUsefulLife if device.remainingUsefulLife is not None else 36.0
                health = int(round(80 + 20 * (rul / 36.0)))
        except Exception as e:
            print(f"Error running pipeline in detail API: {e}")
            
    return {
        "id": device.id,
        "name": device.name,
        "type": device.deviceType,
        "status": device.status,
        "lastUpdated": (latest_telemetry.timestamp if latest_telemetry else device.lastUpdated).isoformat(),
        "battery": battery,
        "temperature": temp,
        "ssd": ssd,
        "cpu": cpu,
        "ram": ram,
        "healthScore": health
    }

@app.get("/predictions/{device_id}", response_model=models.PredictionResponse)
def get_prediction(device_id: str, db: Session = Depends(get_db)):
    # Always attempt to sync latest telemetry from Supabase first
    sync_device_from_supabase(device_id, db)
    
    latest_telemetry = db.query(models.TelemetryRecord)\
        .filter(models.TelemetryRecord.deviceId == device_id)\
        .order_by(models.TelemetryRecord.id.desc())\
        .first()
        
    if not latest_telemetry:
        # Return a default safe prediction if no telemetry exists yet
        return {
            "deviceId": device_id,
            "healthScore": 100,
            "riskLevel": "Low Risk",
            "recommendations": [{
                "title": "Waiting for Data",
                "description": "Waiting for first telemetry reading...",
                "improvement": "N/A",
                "icon": "clock",
                "color": "primary"
            }],
            "shapValues": {},
            "isAnomaly": False,
            "anomalyScore": 0.0,
            "confidenceLevel": 100.0
        }
        
    # 2. Get full telemetry history for cycle estimation
    all_telemetry = db.query(models.TelemetryRecord)\
        .filter(models.TelemetryRecord.deviceId == device_id)\
        .order_by(models.TelemetryRecord.id.asc())\
        .all()
    
    # 3. Run the AI pipeline with history
    telemetry_dict = {
        "deviceId": latest_telemetry.deviceId,
        "cpu": latest_telemetry.cpu,
        "ram": latest_telemetry.ram,
        "battery": latest_telemetry.battery,
        "temperature": latest_telemetry.temperature,
        "ssd": latest_telemetry.ssd
    }
    
    prediction_result = run_ai_pipeline(telemetry_dict, all_telemetry)
    
    # 4. Extract remainingUsefulLife from SQLite DeviceRecord
    device = db.query(models.DeviceRecord).filter(models.DeviceRecord.id == device_id).first()
    rul = device.remainingUsefulLife if (device and device.remainingUsefulLife is not None) else 36.0
    prediction_result["remainingUsefulLife"] = rul
    
    if device and device.deviceType == "phone":
        payload = {}
        import requests
        url = f"https://lonsqhuudhiffjitmcbh.supabase.co/rest/v1/telemetry?device_uuid=eq.{device_id}&order=updated_at.desc&limit=1"
        headers = {
            "apikey": "sb_publishable_huLEhuc-J4bal6hQRkPf5w_O16MKv6V",
            "Authorization": "Bearer sb_publishable_huLEhuc-J4bal6hQRkPf5w_O16MKv6V"
        }
        try:
            r = requests.get(url, headers=headers, timeout=2)
            if r.status_code == 200 and r.json():
                payload = r.json()[0].get("payload", {})
        except Exception:
            pass

        rul = payload.get("native_remaining_useful_life_months", device.remainingUsefulLife if device.remainingUsefulLife is not None else 36.0)
        native_health = payload.get("native_battery_health_percentage", None)
        
        prediction_result["remainingUsefulLife"] = rul
        prediction_result["remaining_useful_life_months"] = rul
        
        def local_safe_float(val, default=0.0):
            if val is None:
                return default
            if isinstance(val, str):
                val = val.replace("%", "").strip()
            try:
                return float(val)
            except:
                return default

        cpu_map = payload.get("cpu") if isinstance(payload.get("cpu"), dict) else {}
        battery_map = payload.get("battery") if isinstance(payload.get("battery"), dict) else {}
        storage_map = payload.get("storage") if isinstance(payload.get("storage"), dict) else {}
        memory_map = payload.get("memory") if isinstance(payload.get("memory"), dict) else {}

        if native_health is not None:
            base_health = float(native_health)
        else:
            # Fallback mapping: RUL (36->0) corresponds to health (100->80)
            base_health = 80.0 + 20.0 * (rul / 36.0)

        # Apply operational deductions to get overall health score
        deductions = 0.0
        
        # 1. Thermal deductions
        temp_val = local_safe_float(cpu_map.get("temperature_c") if cpu_map.get("temperature_c") is not None else cpu_map.get("gpu_temperature_c"), 35.0)
        if temp_val > 38.0:
            deductions += (temp_val - 38.0) * 1.5
        if temp_val > 42.0:
            deductions += (temp_val - 42.0) * 1.0 # additional weight
            
        # 2. Storage deductions
        storage_val = local_safe_float(storage_map.get("usage_percent") if storage_map.get("usage_percent") is not None else storage_map.get("disk_usage_percent"), 0.0)
        if storage_val > 75.0:
            deductions += (storage_val - 75.0) * 0.4
            
        # 3. CPU/RAM load deductions
        cpu_val = local_safe_float(cpu_map.get("usage_percent"), 0.0)
        if cpu_val > 80.0:
            deductions += (cpu_val - 80.0) * 0.2
            
        ram_val = local_safe_float(memory_map.get("usage_percent") if memory_map.get("usage_percent") is not None else memory_map.get("ram_usage_percent"), 0.0)
        if ram_val > 80.0:
            deductions += (ram_val - 80.0) * 0.3
            
        final_health = max(0, min(100, int(round(base_health - deductions))))
        prediction_result["healthScore"] = final_health
            
        # Resource-based risk escalation (Operational Risk)
        battery_level = local_safe_float(battery_map.get("level") if battery_map.get("level") is not None else battery_map.get("percentage"), 100.0)
        storage_usage = local_safe_float(storage_map.get("usage_percent") if storage_map.get("usage_percent") is not None else storage_map.get("disk_usage_percent"), 0.0)
        
        if prediction_result["healthScore"] >= 85:
            if (battery_level < 15 and battery_level > 0) or storage_usage > 75:
                prediction_result["riskLevel"] = "Medium Risk"
            else:
                prediction_result["riskLevel"] = "Low Risk"
        elif prediction_result["healthScore"] >= 75:
            prediction_result["riskLevel"] = "Medium Risk"
        else:
            prediction_result["riskLevel"] = "High Risk"
            
    # 5. Log the prediction run to SQLite database
    db_pred = models.AIPrediction(
        deviceId=device_id,
        healthScore=prediction_result["healthScore"],
        riskLevel=prediction_result["riskLevel"],
        isAnomaly=1 if prediction_result["isAnomaly"] else 0,
        anomalyScore=prediction_result["anomalyScore"],
        confidenceLevel=prediction_result["confidenceLevel"]
    )
    db.add(db_pred)
    db.commit()
    
    return prediction_result

@app.post("/api/support/ticket", response_model=models.SupportTicketResponse)
def create_support_ticket(ticket: models.SupportTicketRequest, db: Session = Depends(get_db)):
    import random
    ticket_id = f"CS-{random.randint(10000, 99999)}"
    
    print(f"[Support Ticket Created] Ticket: {ticket_id} | Device: {ticket.deviceId} | Health: {ticket.healthScore} | Risk: {ticket.riskLevel}")
    print(f"Metrics: CPU={ticket.cpu}%, RAM={ticket.ram}%, Temp={ticket.temperature}C, SSD={ticket.ssd}%")
    
    device = db.query(models.DeviceRecord).filter(models.DeviceRecord.id == ticket.deviceId).first()
    if device:
        device.status = "Support Contacted"
        device.lastUpdated = datetime.utcnow()
        db.commit()
        
    return {
        "status": "success",
        "ticketId": ticket_id,
        "message": f"Support ticket {ticket_id} has been created successfully. Customer support has been notified of your device health issues."
    }

@app.get("/api/reports/{device_id}/docx")
def generate_docx_report(device_id: str, db: Session = Depends(get_db)):
    device = db.query(models.DeviceRecord).filter(models.DeviceRecord.id == device_id).first()
    if not device:
        raise HTTPException(status_code=404, detail="Device not found")
        
    latest_telemetry = db.query(models.TelemetryRecord)\
        .filter(models.TelemetryRecord.deviceId == device_id)\
        .order_by(models.TelemetryRecord.timestamp.desc())\
        .first()
        
    all_telemetry = db.query(models.TelemetryRecord)\
        .filter(models.TelemetryRecord.deviceId == device_id)\
        .order_by(models.TelemetryRecord.timestamp.asc())\
        .all()
        
    if not latest_telemetry:
        raise HTTPException(status_code=400, detail="No telemetry data available for report")
        
    telemetry_dict = {
        "deviceId": latest_telemetry.deviceId,
        "cpu": latest_telemetry.cpu,
        "ram": latest_telemetry.ram,
        "battery": latest_telemetry.battery,
        "temperature": latest_telemetry.temperature,
        "ssd": latest_telemetry.ssd
    }
    prediction = run_ai_pipeline(telemetry_dict, all_telemetry)
    
    doc = Document()
    
    title = doc.add_paragraph()
    title_run = title.add_run("DEVICEGUARDIAN AI DIAGNOSTICS REPORT")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(99, 102, 241)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_meta = doc.add_paragraph()
    p_meta.add_run("Generated At: ").bold = True
    p_meta.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    p_meta.add_run("Device Name: ").bold = True
    p_meta.add_run(f"{device.name}\n")
    p_meta.add_run("Device ID: ").bold = True
    p_meta.add_run(f"{device.id}\n")
    p_meta.add_run("Status: ").bold = True
    p_meta.add_run(f"{device.status}\n")
    
    doc.add_heading("Health & Risk Assessment", level=1)
    p_health = doc.add_paragraph()
    p_health.add_run("AI System Health Score: ").bold = True
    p_health.add_run(f"{prediction['healthScore']}%\n")
    p_health.add_run("Model Confidence Level: ").bold = True
    p_health.add_run(f"{prediction['confidenceLevel']:.1f}%\n")
    p_health.add_run("Risk Classification: ").bold = True
    p_health.add_run(f"{prediction['riskLevel']}\n")
    p_health.add_run("Anomaly Flagged: ").bold = True
    p_health.add_run(f"{'YES' if prediction['isAnomaly'] else 'NO'} (Isolation Forest Score: {prediction['anomalyScore']:.4f})\n")
    
    doc.add_heading("Telemetry Metrics", level=1)
    p_metrics = doc.add_paragraph()
    p_metrics.add_run("CPU Load: ").bold = True
    p_metrics.add_run(f"{latest_telemetry.cpu}%\n")
    p_metrics.add_run("RAM Usage: ").bold = True
    p_metrics.add_run(f"{latest_telemetry.ram}%\n")
    p_metrics.add_run("Battery Level: ").bold = True
    p_metrics.add_run(f"{latest_telemetry.battery}%\n")
    p_metrics.add_run("Core Temperature: ").bold = True
    p_metrics.add_run(f"{latest_telemetry.temperature}°C\n")
    p_metrics.add_run("SSD Storage Used: ").bold = True
    p_metrics.add_run(f"{latest_telemetry.ssd}%\n")

    doc.add_heading("AI Feature Impact (SHAP Analysis)", level=1)
    p_shap = doc.add_paragraph()
    for key, val in prediction['shapValues'].items():
        p_shap.add_run(f"• {key.upper()}: ").bold = True
        p_shap.add_run(f"{val:.4f}\n")
        
    doc.add_heading("Recommended Safe Interventions", level=1)
    for rec in prediction['recommendations']:
        p_rec = doc.add_paragraph()
        p_rec.add_run(f"[{rec['title']}]").bold = True
        p_rec.add_run(f" (Est. Improvement: {rec['improvement']})\n")
        p_rec.add_run(f"{rec['description']}")
        
    filename = f"report_{device_id}.docx"
    filepath = os.path.join(os.getcwd(), filename)
    doc.save(filepath)
    
    return FileResponse(
        path=filepath, 
        filename=f"DeviceGuardian_Report_{device.name.replace(' ', '_')}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
